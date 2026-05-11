import streamlit as st
from streamlit_mic_recorder import mic_recorder
import anthropic
import openai
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import json
import os
import unicodedata
import re
from datetime import datetime


def sanitize_filename(name: str) -> str:
    nfkd = unicodedata.normalize("NFKD", name)
    ascii_str = "".join(c for c in nfkd if not unicodedata.combining(c))
    return re.sub(r"[^\w\-_]", "_", ascii_str)


def get_secret(key: str) -> str:
    """Lee un secreto desde st.secrets o variable de entorno."""
    try:
        return st.secrets[key]
    except (KeyError, FileNotFoundError):
        return os.getenv(key, "")


# ============================================================================
# 1. CONFIGURACIÓN DE PÁGINA
# ============================================================================
st.set_page_config(
    page_title="Voice Draft",
    page_icon="🎙️",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
    <style>
    /* Header principal */
    .vd-header {
        display: flex;
        align-items: center;
        gap: 14px;
        padding: 18px 0 10px 0;
        border-bottom: 2px solid #e8ecf0;
        margin-bottom: 24px;
    }
    .vd-logo {
        font-size: 36px;
        line-height: 1;
    }
    .vd-title {
        font-size: 28px;
        font-weight: 800;
        color: #1a1a2e;
        letter-spacing: -0.5px;
        margin: 0;
        line-height: 1.1;
    }
    .vd-subtitle {
        font-size: 13px;
        color: #6b7280;
        margin: 2px 0 0 0;
    }
    .vd-badge {
        background: linear-gradient(135deg, #0066cc, #0044aa);
        color: white;
        font-size: 10px;
        font-weight: 700;
        padding: 3px 8px;
        border-radius: 20px;
        letter-spacing: 0.5px;
        text-transform: uppercase;
        align-self: flex-start;
        margin-top: 4px;
    }

    /* Contenedores de capas */
    .layer-box {
        border-left: 4px solid #0066cc;
        padding: 10px 14px;
        border-radius: 4px;
        background-color: #f0f4ff;
        margin: 6px 0;
    }
    .layer-complete { border-left-color: #28a745; background-color: #f0fff4; }
    .layer-pending  { border-left-color: #ffc107; background-color: #fffbf0; }
    .layer-error    { border-left-color: #ff4b4b; background-color: #fff5f5; }

    /* Badges */
    .badge {
        display: inline-block;
        padding: 3px 8px;
        border-radius: 12px;
        font-size: 11px;
        font-weight: bold;
        margin-left: 6px;
    }
    .badge-success { background-color: #d4edda; color: #155724; }
    .badge-pending { background-color: #fff3cd; color: #856404; }
    .badge-error   { background-color: #f8d7da; color: #721c24; }

    /* Alerta de secrets */
    .secrets-warning {
        background: #fff8e1;
        border: 1px solid #ffcc02;
        border-radius: 6px;
        padding: 10px 14px;
        font-size: 13px;
        color: #7a5c00;
    }
    </style>
""", unsafe_allow_html=True)

# ============================================================================
# 2. INICIALIZACIÓN DE SESSION STATE
# ============================================================================
def init_session_state():
    if "capas_data" not in st.session_state:
        st.session_state.capas_data = {}
    if "configuracion" not in st.session_state:
        st.session_state.configuracion = {
            "nombre_archivo": "",
            "plantilla": None,
            "modo": "Automático"
        }
    if "historial_procesamiento" not in st.session_state:
        st.session_state.historial_procesamiento = []

init_session_state()

# Cargar API keys desde secrets / env vars
OPENAI_KEY = get_secret("OPENAI_API_KEY")
CLAUDE_KEY  = get_secret("ANTHROPIC_API_KEY")

# ============================================================================
# 3. BARRA LATERAL
# ============================================================================
with st.sidebar:
    st.markdown("### Voice Draft")

    # Indicador de estado de API keys
    oa_ok  = bool(OPENAI_KEY)
    ant_ok = bool(CLAUDE_KEY)
    st.markdown(
        f"{'🟢' if oa_ok  else '🔴'} OpenAI Whisper  \n"
        f"{'🟢' if ant_ok else '🔴'} Anthropic Claude"
    )

    if not oa_ok or not ant_ok:
        st.markdown(
            "<div class='secrets-warning'>"
            "Configura <b>OPENAI_API_KEY</b> y <b>ANTHROPIC_API_KEY</b> "
            "en los Secrets de tu app en Streamlit Cloud."
            "</div>",
            unsafe_allow_html=True
        )

    st.divider()
    st.markdown("##### Plantilla .docx")
    plantilla_subida = st.file_uploader("Sube tu plantilla base", type=["docx"], label_visibility="collapsed")
    if plantilla_subida:
        st.session_state.configuracion["plantilla"] = plantilla_subida
        st.success("Plantilla cargada")

    st.divider()
    st.markdown("##### Estado de Capas")

    if st.session_state.capas_data:
        total_capas    = len(st.session_state.capas_data)
        capas_completas = sum(1 for c in st.session_state.capas_data.values() if c.get("status") == "complete")
        st.progress(capas_completas / total_capas, text=f"{capas_completas}/{total_capas} completadas")

        for nombre_capa, datos in st.session_state.capas_data.items():
            status = datos.get("status", "pending")
            if status == "complete":
                st.success(f"✅ {nombre_capa}", icon=None)
            elif status == "error":
                st.error(f"❌ {nombre_capa}", icon=None)
            else:
                st.info(f"⏳ {nombre_capa}", icon=None)

        st.divider()
        if st.button("🗑️ Limpiar todo", use_container_width=True):
            st.session_state.capas_data = {}
            st.session_state.historial_procesamiento = []
            st.rerun()
    else:
        st.caption("Sin capas grabadas aún.")

# ============================================================================
# 4. HEADER PRINCIPAL
# ============================================================================
st.markdown("""
    <div class="vd-header">
        <div class="vd-logo">🎙️</div>
        <div>
            <p class="vd-title">Voice Draft</p>
            <p class="vd-subtitle">Tu pensamiento en voz, convertido en documento profesional</p>
        </div>
        <div class="vd-badge">by Thinking</div>
    </div>
""", unsafe_allow_html=True)

if not oa_ok or not ant_ok:
    st.warning("Las API keys no están configuradas. La app no podrá procesar audio hasta que estén disponibles.", icon="⚠️")

# ============================================================================
# 5. CONFIGURACIÓN INICIAL
# ============================================================================
st.markdown("#### Paso 1 — Configuración")

col1, col2 = st.columns([0.6, 0.4])
with col1:
    st.session_state.configuracion["nombre_archivo"] = st.text_input(
        "Nombre del archivo final",
        placeholder="Ej: Minuta_Tecnica_2024",
        value=st.session_state.configuracion.get("nombre_archivo", "")
    )
with col2:
    st.session_state.configuracion["modo"] = st.radio(
        "Modo:",
        ["🤖 Automático", "🎯 Manual"],
        horizontal=True
    )

SECCIONES_DISPONIBLES = [
    "Objetivo", "Alcance", "Definiciones", "Descripción",
    "Metodología", "Resultados", "Conclusiones"
]

secciones_a_capturar  = {}
usar_subsecciones_desc = False
num_subsecciones       = 2

if "Manual" in st.session_state.configuracion["modo"]:
    st.markdown("**Secciones a capturar:**")
    cols = st.columns(4)
    for i, seccion in enumerate(SECCIONES_DISPONIBLES):
        with cols[i % 4]:
            secciones_a_capturar[seccion] = st.checkbox(seccion, value=(i < 3))

    if secciones_a_capturar.get("Descripción", False):
        col_check, col_num = st.columns([0.45, 0.55])
        with col_check:
            usar_subsecciones_desc = st.checkbox("Subdividir Descripción", value=False, key="usar_subs")
        if usar_subsecciones_desc:
            with col_num:
                num_subsecciones = st.number_input("¿Cuántas partes?", min_value=2, max_value=5, value=2, key="num_subs")
            secciones_subsecciones = {f"Descripción - Parte {chr(65+i)}": True for i in range(num_subsecciones)}
            if "Descripción" in secciones_a_capturar:
                del secciones_a_capturar["Descripción"]
            secciones_a_capturar.update(secciones_subsecciones)
else:
    secciones_a_capturar = {s: True for s in ["Objetivo", "Alcance", "Definiciones", "Descripción"]}

st.divider()

# ============================================================================
# 6. GRABACIÓN POR CAPAS
# ============================================================================
st.markdown("#### Paso 2 — Grabación de Capas")

capas_pendientes  = [s for s, activa in secciones_a_capturar.items() if activa and s not in st.session_state.capas_data]
capas_completadas = [s for s in st.session_state.capas_data if secciones_a_capturar.get(s, False)]

if capas_pendientes:
    capa_actual = st.selectbox("Sección a grabar:", options=capas_pendientes)
    st.info(f"Grabando: **{capa_actual}**  \n_Espera a que la sección anterior termine de procesarse antes de grabar la siguiente._")

    audio_data = mic_recorder(
        start_prompt="🎤 Iniciar Grabación",
        stop_prompt="🛑 Detener y Procesar",
        key=f"recorder_{capa_actual}"
    )

    if audio_data:
        with st.spinner(f"Procesando '{capa_actual}'..."):
            try:
                # A. Transcripción (Whisper)
                client_oa  = openai.OpenAI(api_key=OPENAI_KEY)
                audio_bio  = io.BytesIO(audio_data["bytes"])
                audio_bio.name = f"audio_{sanitize_filename(capa_actual)}.mp3"

                transcripcion = client_oa.audio.transcriptions.create(
                    model="whisper-1",
                    file=audio_bio
                ).text

                # B. Redacción (Claude)
                client_ant = anthropic.Anthropic(api_key=CLAUDE_KEY)

                prompt = f"""Eres un experto senior en redacción técnica y corporativa.
Transforma la siguiente transcripción informal en una sección profesional para '{capa_actual}'.

REGLAS:
1. NO transcribas literal. Mejora vocabulario con términos técnicos y ejecutivos.
2. Párrafos coherentes y bien hilados (3-5 párrafos).
3. Elimina muletillas (eh, bueno, o sea), repeticiones y errores gramaticales.
4. Tono serio, preciso, nivel corporativo senior.
5. 150-400 palabras según la sección.
6. Devuelve SOLO el texto redactado, sin etiquetas ni markdown.

TRANSCRIPCIÓN:
\"{transcripcion}\"

Redacta ahora la sección '{capa_actual}':"""

                response  = client_ant.messages.create(
                    model="claude-sonnet-4-6",
                    max_tokens=2000,
                    messages=[{"role": "user", "content": prompt}]
                )
                redaccion = response.content[0].text

                st.session_state.capas_data[capa_actual] = {
                    "audio_size":   len(audio_data["bytes"]) / (1024 * 1024),
                    "transcripcion": transcripcion,
                    "redaccion":     redaccion,
                    "status":        "complete",
                    "timestamp":     datetime.now().isoformat()
                }
                st.session_state.historial_procesamiento.append({
                    "capa": capa_actual, "timestamp": datetime.now().isoformat(), "status": "success"
                })
                st.success(f"✅ '{capa_actual}' completada")
                st.rerun()

            except Exception as e:
                st.session_state.capas_data[capa_actual] = {
                    "status": "error", "error_msg": str(e), "timestamp": datetime.now().isoformat()
                }
                st.error(f"Error procesando '{capa_actual}': {str(e)}")
else:
    st.success("✅ Todas las capas han sido grabadas.")

st.divider()

# ============================================================================
# 7. REVISIÓN DE CAPAS
# ============================================================================
if st.session_state.capas_data:
    st.markdown("#### Paso 3 — Revisión de Capas")

    for nombre_capa, datos in st.session_state.capas_data.items():
        status = datos.get("status", "pending")

        css_class = {"complete": "layer-complete", "error": "layer-error"}.get(status, "layer-pending")
        icon      = {"complete": "✅", "error": "❌"}.get(status, "⏳")
        badge_cls = {"complete": "badge-success", "error": "badge-error"}.get(status, "badge-pending")
        badge_txt = {"complete": "Completada", "error": "Error"}.get(status, "Pendiente")

        container = st.container(border=True)
        container.markdown(
            f"<div class='layer-box {css_class}'>"
            f"<strong>{icon} {nombre_capa}</strong>"
            f"<span class='badge {badge_cls}'>{badge_txt}</span>"
            f"</div>",
            unsafe_allow_html=True
        )

        with container:
            if status == "complete":
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.caption(f"Audio: {datos.get('audio_size', 0):.2f} MB")
                with col2:
                    st.caption(datos.get("timestamp", "")[:10])
                with col3:
                    if st.button("🗑️ Eliminar", key=f"del_{nombre_capa}"):
                        del st.session_state.capas_data[nombre_capa]
                        st.rerun()

                tab1, tab2 = st.tabs(["Transcripción original", "Redacción pulida"])
                with tab1:
                    st.text_area("", value=datos.get("transcripcion", ""), height=110,
                                 disabled=True, key=f"trans_{nombre_capa}", label_visibility="collapsed")
                with tab2:
                    redaccion_editada = st.text_area("", value=datos.get("redaccion", ""),
                                                     height=140, key=f"redac_{nombre_capa}", label_visibility="collapsed")
                    if redaccion_editada != datos.get("redaccion", ""):
                        st.session_state.capas_data[nombre_capa]["redaccion"] = redaccion_editada
                        st.caption("✏️ Guardado")

            elif status == "error":
                st.error(datos.get("error_msg", "Error desconocido"))
                if st.button("🔄 Reintentar", key=f"retry_{nombre_capa}"):
                    del st.session_state.capas_data[nombre_capa]
                    st.rerun()

    st.divider()

# ============================================================================
# 8. ENSAMBLADOR
# ============================================================================
st.markdown("#### Paso 4 — Ensamblar Documento")

capas_ok = st.session_state.capas_data and all(
    d.get("status") == "complete" for d in st.session_state.capas_data.values()
)

if capas_ok:
    col_btn, col_info = st.columns([0.3, 0.7])
    with col_btn:
        ensamblar = st.button("🚀 Ensamblar", type="primary", use_container_width=True)
    with col_info:
        nombre = st.session_state.configuracion.get("nombre_archivo", "documento")
        st.markdown(f"**{nombre}.docx** — {len(st.session_state.capas_data)} secciones listas")

    if ensamblar:
        with st.spinner("Ensamblando documento..."):
            try:
                doc = (Document(st.session_state.configuracion["plantilla"])
                       if st.session_state.configuracion.get("plantilla") else Document())

                if not doc.paragraphs:
                    doc.add_heading(st.session_state.configuracion.get("nombre_archivo", "Documento"), level=1)

                TAGS_MAP = {
                    "Objetivo":    "{{OBJETIVO}}",
                    "Alcance":     "{{ALCANCE}}",
                    "Definiciones":"{{DEFINICIONES}}",
                    "Descripción": "{{DESCRIPCION}}",
                    "Metodología": "{{METODOLOGIA}}",
                    "Resultados":  "{{RESULTADOS}}",
                    "Conclusiones":"{{CONCLUSIONES}}"
                }

                def reemplazar_tag(doc, tag, contenido):
                    reemplazado = False
                    for p in doc.paragraphs:
                        if tag in p.text:
                            p.text = p.text.replace(tag, contenido)
                            reemplazado = True
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    if tag in p.text:
                                        p.text = p.text.replace(tag, contenido)
                                        reemplazado = True
                    if not reemplazado:
                        nombres = [k for k, v in TAGS_MAP.items() if v == tag]
                        if nombres:
                            doc.add_heading(nombres[0], level=2)
                            doc.add_paragraph(contenido)

                desc_subs  = [(n, d.get("redaccion", "")) for n, d in st.session_state.capas_data.items() if "Descripción - Parte" in n]
                otras_capas = [(n, d.get("redaccion", "")) for n, d in st.session_state.capas_data.items() if "Descripción - Parte" not in n]

                if desc_subs:
                    contenido_desc = "".join(f"{n}\n\n{r}\n\n" for n, r in desc_subs)
                    reemplazar_tag(doc, "{{DESCRIPCION}}", contenido_desc)

                for nombre_capa, redaccion in otras_capas:
                    if nombre_capa in TAGS_MAP:
                        reemplazar_tag(doc, TAGS_MAP[nombre_capa], redaccion)

                bio = io.BytesIO()
                doc.save(bio)
                bio.seek(0)

                st.success("✅ Documento ensamblado")

                with st.expander("Vista previa"):
                    for nc in st.session_state.capas_data:
                        r = st.session_state.capas_data[nc].get("redaccion", "")
                        st.markdown(f"**{nc}**")
                        st.write(r[:200] + "..." if len(r) > 200 else r)

                archivo_nombre = st.session_state.configuracion.get("nombre_archivo", "documento")
                st.download_button(
                    label=f"📥 Descargar {archivo_nombre}.docx",
                    data=bio.getvalue(),
                    file_name=f"{archivo_nombre}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

                st.session_state.historial_procesamiento.append({
                    "evento": "documento_ensamblado",
                    "nombre_archivo": archivo_nombre,
                    "num_capas": len(st.session_state.capas_data),
                    "timestamp": datetime.now().isoformat()
                })

            except Exception as e:
                st.error(f"Error ensamblando: {str(e)}")
else:
    pendientes = sum(1 for d in st.session_state.capas_data.values() if d.get("status") != "complete")
    if pendientes > 0:
        st.warning(f"Completa las {pendientes} capas pendientes para ensamblar.")
    else:
        st.info("Graba al menos una capa para habilitar el ensamblador.")

st.divider()

# ============================================================================
# 9. HISTORIAL
# ============================================================================
with st.expander("Historial de sesión"):
    if st.session_state.historial_procesamiento:
        st.json(st.session_state.historial_procesamiento)
    else:
        st.caption("Sin historial aún.")
    if st.button("📥 Exportar historial JSON"):
        historial_json = json.dumps(st.session_state.historial_procesamiento, indent=2, ensure_ascii=False)
        st.download_button("Descargar", data=historial_json,
                           file_name="historial.json", mime="application/json")
